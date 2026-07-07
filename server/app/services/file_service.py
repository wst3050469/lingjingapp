"""灵境 - 文件处理服务（存储、文本提取、图片AI描述、视频处理）"""
import os
import uuid
import base64
import asyncio
import json
import logging
import subprocess
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)

import aiofiles
import httpx
from PIL import Image

import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import db as database
from app.config import (
    UPLOAD_DIR, VISION_MODEL, VISION_TIMEOUT,
    VISION_FALLBACK_MODELS, VISION_REMOTE_MODELS,
    MAX_EXTRACT_CHARS, OLLAMA_PRIMARY, OLLAMA_SECONDARY,
    MAX_DRAWING_SIZE,
)

# 允许的文件类型
ALLOWED_IMAGE_EXT = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'}
ALLOWED_DOC_EXT = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.txt'}
ALLOWED_DRAWING_EXT = {'.dwg', '.dxf'}
ALLOWED_VIDEO_EXT = {'.mp4', '.mov', '.avi', '.mkv', '.webm', '.3gp'}
ALLOWED_EXTENSIONS = ALLOWED_IMAGE_EXT | ALLOWED_DOC_EXT | ALLOWED_DRAWING_EXT | ALLOWED_VIDEO_EXT

# OS级最大文件大小（超出将跳过AI分析）
_MAX_AI_ANALYSIS_SIZE = 15 * 1024 * 1024  # 15MB——超过此大小的文件只保存不上传AI分析

# 并发控制：同时最多2个图片分析任务
_vision_semaphore = asyncio.Semaphore(2)

# 视觉模型健康状态
_vision_model_healthy = True
_vision_model_lock = asyncio.Lock()
_vision_check_fail_count = 0  # 连续失败计数，用于抑制重复日志
_VISION_CHECK_LOG_INTERVAL = 6  # 每6次失败（即每30分钟）才记录一次警告


async def _ensure_vision_model() -> bool:
    """检查并确保视觉模型可用，不可用时自动重启"""
    global _vision_model_healthy, _vision_check_fail_count
    async with _vision_model_lock:
        try:
            async with httpx.AsyncClient(timeout=10) as client:
                resp = await client.post(
                    f"{OLLAMA_PRIMARY}/api/generate",
                    json={"model": VISION_MODEL, "prompt": "ping", "stream": False},
                )
                if resp.status_code == 200:
                    _vision_model_healthy = True
                    _vision_check_fail_count = 0
                    return True
                elif "not found" in resp.text.lower():
                    logger.warning(f"视觉模型 {VISION_MODEL} 不存在，尝试拉取...")
                    pull_resp = await client.post(
                        f"{OLLAMA_PRIMARY}/api/pull",
                        json={"name": VISION_MODEL},
                    )
                    if pull_resp.status_code == 200:
                        _vision_model_healthy = True
                        _vision_check_fail_count = 0
                        return True
                _vision_model_healthy = False
                return False
        except Exception as e:
            _vision_model_healthy = False
            _vision_check_fail_count += 1
            # 首次失败或每隔N次失败才记录警告，避免日志刷屏
            if _vision_check_fail_count == 1:
                logger.warning(f"视觉模型健康检查失败（Ollama不可达，图片分析将不可用）: {e}")
            elif _vision_check_fail_count % _VISION_CHECK_LOG_INTERVAL == 0:
                logger.warning(
                    f"视觉模型仍不可达（已连续失败{_vision_check_fail_count}次，" +
                    f"约{_vision_check_fail_count * 5}分钟），请检查Ollama服务: {e}"
                )
            return False


async def _restart_vision_model():
    """重启视觉模型（Ollama模型卸载后重新加载）"""
    global _vision_model_healthy
    try:
        logger.info(f"正在重启视觉模型 {VISION_MODEL} ...")
        async with httpx.AsyncClient(timeout=30) as client:
            # 先卸载
            await client.delete(
                f"{OLLAMA_PRIMARY}/api/pull",
                json={"name": VISION_MODEL},
            )
            await asyncio.sleep(1)
            # 重新加载（通过generate触发）
            resp = await client.post(
                f"{OLLAMA_PRIMARY}/api/generate",
                json={"model": VISION_MODEL, "prompt": "ping", "stream": False, "keep_alive": "5m"},
            )
            if resp.status_code == 200:
                _vision_model_healthy = True
                logger.info(f"视觉模型 {VISION_MODEL} 重启成功")
                return True
            else:
                logger.error(f"视觉模型重启失败: {resp.status_code} {resp.text[:200]}")
                _vision_model_healthy = False
                return False
    except Exception as e:
        logger.error(f"视觉模型重启异常: {e}")
        _vision_model_healthy = False
        return False


async def _vision_health_check_loop():
    """后台定时检查视觉模型健康状态（每5分钟）"""
    while True:
        await _ensure_vision_model()
        await asyncio.sleep(300)  # 5分钟


def classify_file(filename: str) -> str | None:
    """根据扩展名判断文件类型，返回 image/pdf/word/excel/txt/drawing/video 或 None"""
    ext = Path(filename).suffix.lower()
    if ext in ALLOWED_IMAGE_EXT:
        return "image"
    if ext == '.pdf':
        return "pdf"
    if ext in ('.doc', '.docx'):
        return "word"
    if ext in ('.xls', '.xlsx'):
        return "excel"
    if ext == '.txt':
        return "txt"
    if ext in ALLOWED_DRAWING_EXT:
        return "drawing"
    if ext in ALLOWED_VIDEO_EXT:
        return "video"
    return None


async def save_upload(file_bytes: bytes, filename: str, session_id: str,
                      invite_code: str, mime_type: str | None = None) -> dict:
    """保存上传文件到磁盘并写入DB，返回文件元信息"""
    file_type = classify_file(filename)
    if not file_type:
        raise ValueError(f"不支持的文件类型: {filename}")

    ext = Path(filename).suffix.lower()
    file_id = f"f_{uuid.uuid4().hex[:16]}"
    month_dir = datetime.now().strftime("%Y-%m")
    rel_dir = os.path.join(invite_code, month_dir)
    abs_dir = os.path.join(UPLOAD_DIR, rel_dir)
    os.makedirs(abs_dir, exist_ok=True)

    stored_name = f"{file_id}{ext}"
    stored_path = os.path.join(abs_dir, stored_name)
    url = f"/uploads/{rel_dir}/{stored_name}"

    # 异步写入文件
    async with aiofiles.open(stored_path, 'wb') as f:
        await f.write(file_bytes)

    # 图片生成缩略图
    thumbnail_path = None
    thumbnail_url = None
    if file_type == "image":
        try:
            thumb_name = f"{file_id}_thumb.jpg"
            thumb_path = os.path.join(abs_dir, thumb_name)
            img = Image.open(stored_path)
            img.thumbnail((300, 300))
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            img.save(thumb_path, "JPEG", quality=80)
            thumbnail_path = thumb_path
            thumbnail_url = f"/uploads/{rel_dir}/{thumb_name}"
        except Exception as e:
            logging.getLogger("lingjing.file_service").warning(f"缩略图生成失败: {e}")
            pass

    # 写入数据库
    async with database.pool.acquire() as conn:
        await conn.execute(
            """INSERT INTO message_files
               (file_id, session_id, invite_code, original_name, stored_path,
                file_type, mime_type, file_size, thumbnail_path, status)
               VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, 'uploaded')""",
            file_id, session_id, invite_code, filename, stored_path,
            file_type, mime_type, len(file_bytes), thumbnail_path,
        )

    return {
        "file_id": file_id,
        "type": file_type,
        "name": filename,
        "size": len(file_bytes),
        "url": url,
        "thumbnail_url": thumbnail_url,
    }


async def process_file(file_id: str):
    """后台异步处理文件：提取文本或AI描述图片"""
    async with database.pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT file_id, stored_path, file_type FROM message_files WHERE file_id=$1",
            file_id,
        )
    if not row:
        return

    stored_path = row["stored_path"]
    file_type = row["file_type"]

    # 更新状态为处理中
    async with database.pool.acquire() as conn:
        await conn.execute(
            "UPDATE message_files SET status='processing' WHERE file_id=$1", file_id)

    extracted_text = None
    ai_description = None

    try:
        if file_type == "image":
            ai_description = await _describe_image(stored_path)
            # 单据识别（后台异步，不阻塞）
            if ai_description and "[图片分析" not in ai_description:
                try:
                    from .receipt_service import analyze_receipt
                    receipt_info = await analyze_receipt(ai_description)
                    if receipt_info and receipt_info.get("is_receipt"):
                        async with database.pool.acquire() as conn2:
                            await conn2.execute(
                                "UPDATE message_files SET ai_tags=$2 WHERE file_id=$1",
                                file_id, json.dumps(receipt_info, ensure_ascii=False),
                            )
                except Exception:
                    logger.warning("更新AI识别标签到message_files失败", exc_info=True)
                    pass
        elif file_type == "pdf":
            extracted_text = await asyncio.to_thread(_extract_pdf, stored_path)
        elif file_type == "word":
            if stored_path.lower().endswith('.docx'):
                extracted_text = await asyncio.to_thread(_extract_docx, stored_path)
            elif stored_path.lower().endswith('.doc'):
                extracted_text = await asyncio.to_thread(_extract_doc, stored_path)
        elif file_type == "excel":
            extracted_text = await asyncio.to_thread(_extract_xlsx, stored_path)
        elif file_type == "txt":
            extracted_text = await asyncio.to_thread(_extract_txt, stored_path)
        elif file_type == "drawing":
            # 图纸文件：尝试 DXF 解析 → AI 视觉分析描述（如果有预览图）
            extracted_text = await asyncio.to_thread(_analyze_drawing, stored_path, filename)
            if not extracted_text or len(extracted_text) < 10:
                # 如果 DXF 解析为空，尝试用 AI 视觉分析（对扫描版/截图有效）
                file_size = os.path.getsize(stored_path)
                if file_size < _MAX_AI_ANALYSIS_SIZE:
                    ai_description = await _describe_image(stored_path)
                    extracted_text = ai_description or "[图纸分析：无法解析此文件格式]"
                else:
                    extracted_text = f"[图纸文件过大({file_size//1024//1024}MB)，跳过AI分析]"

        status = "ready"
    except Exception as e:
        extracted_text = f"[文件处理失败: {str(e)[:100]}]"
        status = "error"

    async with database.pool.acquire() as conn:
        await conn.execute(
            """UPDATE message_files
               SET extracted_text=$2, ai_description=$3, status=$4
               WHERE file_id=$1""",
            file_id, extracted_text, ai_description, status,
        )


def _extract_pdf(path: str) -> str:
    """PyMuPDF提取PDF文本"""
    import fitz
    doc = fitz.open(path)
    texts = []
    total = 0
    for page in doc:
        text = page.get_text().strip()
        if text:
            texts.append(text)
            total += len(text)
            if total >= MAX_EXTRACT_CHARS:
                break
    doc.close()
    result = "\n\n".join(texts)
    return result[:MAX_EXTRACT_CHARS]


def _extract_docx(path: str) -> str:
    """python-docx提取Word文本"""
    from docx import Document
    doc = Document(path)
    texts = []
    total = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
            total += len(text)
            if total >= MAX_EXTRACT_CHARS:
                break
    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                texts.append(" | ".join(cells))
                total += len(texts[-1])
                if total >= MAX_EXTRACT_CHARS:
                    break
    result = "\n".join(texts)
    return result[:MAX_EXTRACT_CHARS]


def _extract_doc(path: str) -> str:
    """olefile提取旧版 .doc (Word 97-2003) 文本"""
    path_str = str(path)
    texts = []
    total = 0
    try:
        # 方法1: 使用olefile读取WordDocument流中的文本
        import olefile
        if olefile.isOleFile(path_str):
            ole = olefile.OleFileIO(path_str)
            # 尝试读取 WordDocument 流的文本
            if ole.exists('WordDocument'):
                data = ole.openstream('WordDocument').read()
                # 简单提取可读的ASCII/Unicode文本
                # DOC格式中文本以UTF-16LE编码存储在流中
                try:
                    raw = data.decode('utf-16-le', errors='ignore')
                    # 过滤控制字符，保留可读文本
                    import re
                    readable = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', raw)
                    for line in readable.split('\n'):
                        line = line.strip()
                        if line and len(line) > 1:
                            texts.append(line)
                            total += len(line)
                            if total >= MAX_EXTRACT_CHARS:
                                break
                except Exception:
                    pass
            ole.close()
    except Exception:
        pass

    if texts:
        return "\n".join(texts)[:MAX_EXTRACT_CHARS]
    return f"[Word 97-2003 (.doc) 格式文件，已保存但文本提取受限，建议另存为 .docx 后重新上传]"


def _extract_xlsx(path: str) -> str:
    """openpyxl提取Excel文本"""
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    texts = []
    total = 0
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        texts.append(f"[工作表: {sheet_name}]")
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = " | ".join(cells)
            if line.strip(" |"):
                texts.append(line)
                total += len(line)
                row_count += 1
                if total >= MAX_EXTRACT_CHARS or row_count > 200:
                    break
        if total >= MAX_EXTRACT_CHARS:
            break
    wb.close()
    result = "\n".join(texts)
    return result[:MAX_EXTRACT_CHARS]


def _extract_txt(path: str) -> str:
    """提取TXT纯文本内容"""
    try:
        # 自动检测编码（优先UTF-8，失败后回退GBK/GB2312）
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1', 'utf-16']
        for enc in encodings:
            try:
                with open(path, 'r', encoding=enc, errors='ignore') as f:
                    text = f.read(MAX_EXTRACT_CHARS + 500)
                    return text[:MAX_EXTRACT_CHARS]
            except (UnicodeDecodeError, LookupError):
                continue
        # 全部失败，以二进制读取
        with open(path, 'rb') as f:
            raw = f.read(MAX_EXTRACT_CHARS)
            return raw.decode('utf-8', errors='replace')[:MAX_EXTRACT_CHARS]
    except Exception as e:
        logger.warning(f"TXT提取失败: {e}")
        return f"[TXT文件读取失败: {str(e)[:100]}]"


def _analyze_drawing(path: str, filename: str) -> str:
    """分析图纸文件（DXF/DWG），提取结构信息

    策略：
    1. DXF 文件 → 使用 ezdxf 提取图层、文本、尺寸标注等信息
    2. DWG 文件 → 返回文件元信息（无法直接解析私有格式）
    """
    ext = Path(filename).suffix.lower()
    file_size = os.path.getsize(path)
    size_mb = file_size / (1024 * 1024)

    parts = [f"[图纸文件] 名称: {filename}", f"大小: {size_mb:.1f}MB"]

    if ext == '.dxf':
        try:
            import ezdxf
            doc = ezdxf.readfile(path)

            # 提取图层信息
            layers = doc.layers
            layer_names = [layer.dxf.name for layer in layers if hasattr(layer, 'dxf')]
            if layer_names:
                parts.append(f"图层({len(layer_names)}): {', '.join(layer_names[:30])}")

            # 提取文本内容（TEXT/MTEXT实体）
            texts = []
            msp = doc.modelspace()
            for entity in msp:
                dxf_type = entity.dxftype()
                if dxf_type == 'TEXT':
                    texts.append(entity.dxf.text)
                elif dxf_type == 'MTEXT':
                    texts.append(entity.dxf.text)
                elif dxf_type == 'DIMENSION':
                    try:
                        dim_text = entity.get_measurement()
                        texts.append(f"尺寸: {dim_text:.2f}")
                    except Exception:
                        pass

            if texts:
                text_content = '; '.join(texts)
                if len(text_content) > 500:
                    text_content = text_content[:500] + '...'
                parts.append(f"标注/文字内容: {text_content}")
            else:
                parts.append("标注/文字: 无文本标注信息")

            # 提取实体数量统计
            entity_counts = {}
            for entity in msp:
                t = entity.dxftype()
                entity_counts[t] = entity_counts.get(t, 0) + 1

            if entity_counts:
                counts_str = ', '.join(f"{k}:{v}" for k, v in sorted(entity_counts.items(), key=lambda x: -x[1])[:10])
                parts.append(f"图元统计: {counts_str}")

            doc.close()

        except ImportError:
            parts.append("解析: DXF解析库未安装(ezdxf)")
        except Exception as e:
            logger.warning(f"DXF解析失败 {path}: {e}")
            parts.append(f"解析: DXF格式解析异常 ({str(e)[:80]})")
    else:
        # DWG (AutoCAD专有二进制格式)
        parts.append("格式: AutoCAD DWG (二进制专有格式)")
        parts.append("提示: 如需查看内容，请将图纸导出为PDF或截图后重新上传")

    return "\n".join(parts)


async def _describe_image(path: str) -> str:
    """调用Ollama多模态模型描述图片"""

    async with _vision_semaphore:
        # 读取并base64编码图片
        async with aiofiles.open(path, 'rb') as f:
            img_bytes = await f.read()
        img_b64 = base64.b64encode(img_bytes).decode('utf-8')
        file_size = len(img_bytes)
        logger.info(f"图片分析: {path} ({file_size} bytes)")

        prompt = "请仔细识别这张图片中的所有文字内容，包括数字、字母和中文。同时描述图片的场景和关键信息。用中文回答，控制在500字以内。"

        # 构建目标列表：(url, model_name)
        targets = []
        # 本地主模型
        targets.append((OLLAMA_PRIMARY, VISION_MODEL))
        # 本地备用模型
        for m in VISION_FALLBACK_MODELS:
            targets.append((OLLAMA_PRIMARY, m))
        # 远程模型
        for m in VISION_REMOTE_MODELS:
            targets.append((OLLAMA_SECONDARY, m))

        for url, model in targets:
            try:
                # 检查模型健康状态，不健康时自动重启
                if not _vision_model_healthy:
                    logger.info("视觉模型状态异常，尝试重启...")
                    await _restart_vision_model()

                payload = {
                    "model": model,
                    "messages": [{
                        "role": "user",
                        "content": prompt,
                        "images": [img_b64],
                    }],
                    "stream": False,
                }
                logger.info(f"尝试 {url} 模型 {model} ...")
                async with httpx.AsyncClient(timeout=VISION_TIMEOUT) as client:
                    resp = await client.post(f"{url}/api/chat", json=payload)
                if resp.status_code == 200:
                    data = resp.json()
                    content = data.get("message", {}).get("content", "")
                    if content:
                        logger.info(f"图片分析成功 ({model}): {content[:80]}...")
                        _vision_model_healthy = True
                        return content
                    logger.warning(f"图片分析返回空内容 ({model})")
                elif "model runner has unexpectedly stopped" in resp.text:
                    logger.warning(f"视觉模型 {model} 崩溃，尝试自动重启...")
                    _vision_model_healthy = False
                    await _restart_vision_model()
                    # 重启后重试一次
                    continue
                else:
                    logger.warning(f"图片分析失败 ({model}): HTTP {resp.status_code} {resp.text[:200]}")
            except Exception as e:
                logger.warning(f"图片分析异常 ({model} at {url}): {type(e).__name__}: {e}")

        logger.error(f"图片分析全部失败: {path}")
        return "[图片分析暂时不可用，请稍后重试]"


async def get_file_contexts(file_ids: list[str]) -> list[dict]:
    """获取多个文件的上下文信息，用于注入AI对话"""
    if not file_ids:
        return []

    results = []
    async with database.pool.acquire() as conn:
        for fid in file_ids:
            row = await conn.fetchrow(
                """SELECT file_id, original_name, file_type, file_size,
                          extracted_text, ai_description, status,
                          stored_path, thumbnail_path, ai_tags
                   FROM message_files WHERE file_id=$1""",
                fid,
            )
            if not row:
                continue

            context_text = ""
            if row["file_type"] == "image":
                desc = row["ai_description"] or "[图片分析中...]"
                context_text = desc
            else:
                text = row["extracted_text"] or "[文档内容提取中...]"
                context_text = text

            # 构造URL — OSS 文件用 CDN，本地文件用 /uploads
            stored_path = row["stored_path"]
            if stored_path.startswith("uploads/"):
                from app.services.oss_service import get_cdn_url
                url = get_cdn_url(stored_path)
            else:
                rel = stored_path.replace(UPLOAD_DIR, "").lstrip("/")
                url = f"/uploads/{rel}"

            thumb_url = None
            if row["thumbnail_path"]:
                thumb_rel = row["thumbnail_path"].replace(UPLOAD_DIR, "").lstrip("/")
                thumb_url = f"/uploads/{thumb_rel}"

            results.append({
                "file_id": row["file_id"],
                "type": row["file_type"],
                "name": row["original_name"],
                "size": row["file_size"],
                "url": url,
                "thumbnail_url": thumb_url,
                "context_text": context_text,
                "status": row["status"],
                "ai_tags": row.get("ai_tags") or {},
            })

    return results


async def wait_for_processing(file_ids: list[str], timeout: int = 30):
    """等待文件处理完成，最多等timeout秒"""
    if not file_ids:
        return
    for _ in range(timeout):
        async with database.pool.acquire() as conn:
            count = await conn.fetchval(
                """SELECT count(*) FROM message_files
                   WHERE file_id = ANY($1) AND status IN ('uploaded', 'processing')""",
                file_ids,
            )
        if count == 0:
            return
        await asyncio.sleep(1)


# ============================================================
# OSS 文件处理（从 OSS 下载 → 本地处理 → 更新 DB）
# ============================================================

async def process_oss_file(file_id: str, object_key: str, file_type: str):
    """从 OSS 下载文件并异步处理（图片AI描述/文档文本提取/视频缩略图）"""

    # 优先使用本地 _pending_oss 目录的文件（避免CDN下载延迟）
    local_path = None
    pending_path = os.path.join(UPLOAD_DIR, "_pending_oss", os.path.basename(object_key))
    if os.path.exists(pending_path):
        local_path = pending_path
        logger.info(f"使用本地缓存文件: {pending_path}")
    
    # 本地没有则从 CDN 下载到临时目录
    if not local_path:
        from app.services.oss_service import download_to_temp
        local_path = await download_to_temp(object_key)
        if not local_path:
            # CDN下载失败，尝试直接从OSS下载
            from app.services.oss_service import download_from_oss
            local_path = await download_from_oss(object_key)

    if not local_path:
        logger.error(f"OSS 文件下载失败: {object_key}")
        # 处理失败但OSS文件完整，标记为ready让用户可访问
        async with database.pool.acquire() as conn:
            await conn.execute(
                "UPDATE message_files SET status='ready', extracted_text='[AI分析暂不可用]' WHERE file_id=$1",
                file_id,
            )
        return

    extracted_text = None
    ai_description = None
    thumbnail_path = None

    # 确保文件存在（可能被并发清理）
    if local_path and not os.path.exists(local_path):
        logger.warning(f"本地文件已被清理，重新下载: {local_path}")
        from app.services.oss_service import download_to_temp
        local_path = await download_to_temp(object_key)
        if not local_path:
            async with database.pool.acquire() as conn:
                await conn.execute(
                    "UPDATE message_files SET status='ready', extracted_text='[文件下载失败，请重试]' WHERE file_id=$1",
                    file_id,
                )
            return

    try:
        if file_type == "image":
            ai_description = await _describe_image(local_path)
            # 单据识别（后台异步）
            if ai_description and "[图片分析" not in ai_description:
                try:
                    from .receipt_service import analyze_receipt
                    receipt_info = await analyze_receipt(ai_description)
                    if receipt_info and receipt_info.get("is_receipt"):
                        async with database.pool.acquire() as conn2:
                            await conn2.execute(
                                "UPDATE message_files SET ai_tags=$2 WHERE file_id=$1",
                                file_id, json.dumps(receipt_info, ensure_ascii=False),
                            )
                except Exception:
                    logger.warning("更新AI标签失败(file_service)", exc_info=True)
                    pass
            # 生成缩略图
            try:
                thumb_dir = os.path.join(UPLOAD_DIR, "thumbnails")
                os.makedirs(thumb_dir, exist_ok=True)
                thumb_name = f"{file_id}_thumb.jpg"
                thumb_full = os.path.join(thumb_dir, thumb_name)
                img = Image.open(local_path)
                img.thumbnail((300, 300))
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                img.save(thumb_full, "JPEG", quality=80)
                thumbnail_path = f"/uploads/thumbnails/{thumb_name}"
            except Exception as e:
                logger.warning(f"缩略图生成失败: {e}")

        elif file_type == "video":
            # 视频：提取元信息 + 截帧缩略图
            video_meta = await asyncio.to_thread(_extract_video_meta, local_path)
            extracted_text = json.dumps(video_meta, ensure_ascii=False) if video_meta else "[视频信息提取失败]"

            # 截取第 1 帧作为封面缩略图
            try:
                thumb_dir = os.path.join(UPLOAD_DIR, "thumbnails")
                os.makedirs(thumb_dir, exist_ok=True)
                thumb_name = f"{file_id}_thumb.jpg"
                thumb_full = os.path.join(thumb_dir, thumb_name)
                await asyncio.to_thread(_extract_video_thumbnail, local_path, thumb_full)
                if os.path.exists(thumb_full) and os.path.getsize(thumb_full) > 0:
                    thumbnail_path = f"/uploads/thumbnails/{thumb_name}"
            except Exception as e:
                logger.warning(f"视频缩略图生成失败: {e}")

        elif file_type == "pdf":
            extracted_text = await asyncio.to_thread(_extract_pdf, local_path)
        elif file_type == "word":
            if local_path.lower().endswith('.docx'):
                extracted_text = await asyncio.to_thread(_extract_docx, local_path)
            elif local_path.lower().endswith('.doc'):
                extracted_text = await asyncio.to_thread(_extract_doc, local_path)
        elif file_type == "excel":
            extracted_text = await asyncio.to_thread(_extract_xlsx, local_path)
        elif file_type == "txt":
            extracted_text = await asyncio.to_thread(_extract_txt, local_path)
        elif file_type == "drawing":
            # 图纸文件（OSS来源）
            extracted_text = await asyncio.to_thread(_analyze_drawing, local_path, os.path.basename(local_path))
            if (not extracted_text or len(extracted_text) < 10) and file_size < _MAX_AI_ANALYSIS_SIZE:
                ai_description = await _describe_image(local_path)
                extracted_text = ai_description or "[图纸分析：无法解析此文件格式]"

        status = "ready"
    except Exception as e:
        logger.error(f"文件处理失败 {file_id}: {e}")
        extracted_text = f"[文件处理失败: {str(e)[:100]}]"
        status = "error"
    finally:
        # 清理临时文件
        try:
            os.remove(local_path)
        except Exception:
            logger.warning(f"清理临时文件失败: {local_path}", exc_info=True)
            pass

    # 更新 DB
    async with database.pool.acquire() as conn:
        await conn.execute(
            """UPDATE message_files
               SET extracted_text=$2, ai_description=$3, thumbnail_path=$4, status=$5
               WHERE file_id=$1""",
            file_id, extracted_text, ai_description, thumbnail_path, status,
        )

    logger.info(f"OSS 文件处理完成: {file_id} type={file_type} status={status}")


def _extract_video_meta(filepath: str) -> dict | None:
    """用 ffprobe 提取视频元信息：时长、分辨率、编码"""
    try:
        result = subprocess.run(
            [
                "ffprobe", "-v", "quiet", "-print_format", "json",
                "-show_format", "-show_streams", filepath,
            ],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            return None

        data = json.loads(result.stdout)
        fmt = data.get("format", {})
        streams = data.get("streams", [])

        video_stream = None
        for s in streams:
            if s.get("codec_type") == "video":
                video_stream = s
                break

        duration = float(fmt.get("duration", 0))
        size = int(fmt.get("size", 0))

        return {
            "duration": round(duration, 1),
            "duration_str": f"{int(duration // 60)}分{int(duration % 60)}秒",
            "size_mb": round(size / (1024 * 1024), 1),
            "resolution": f"{video_stream.get('width')}x{video_stream.get('height')}" if video_stream else None,
            "codec": video_stream.get("codec_name") if video_stream else None,
            "bitrate": fmt.get("bit_rate"),
        }
    except FileNotFoundError:
        # ffprobe 未安装
        return {"duration": 0, "error": "ffprobe 未安装"}
    except Exception as e:
        return {"error": str(e)[:100]}


def _extract_video_thumbnail(filepath: str, output_path: str):
    """用 ffmpeg 截取视频第 1 秒帧作为缩略图"""
    try:
        subprocess.run(
            [
                "ffmpeg", "-y", "-ss", "00:00:01", "-i", filepath,
                "-vframes", "1", "-q:v", "3",
                "-vf", "scale=640:-1",
                output_path,
            ],
            capture_output=True, timeout=30,
        )
    except FileNotFoundError:
        pass  # ffmpeg 未安装
    except Exception:
        logger.warning("ffmpeg处理异常", exc_info=True)
        pass


# 模块启动时启动视觉模型健康检查后台任务
def start_vision_health_check():
    """启动视觉模型健康检查后台循环"""
    task = asyncio.create_task(_vision_health_check_loop())
    logger.info("视觉模型健康检查后台任务已启动")
    return task
